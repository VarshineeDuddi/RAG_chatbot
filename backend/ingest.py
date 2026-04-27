"""
ingest.py — RAG Steps 1-3
- Better embeddings: BAAI/bge-base-en-v1.5 (stronger than all-MiniLM-L6-v2)
- Supports: PDF (.pdf) and Word (.docx)
"""

import os
import uuid
import io
import pymupdf
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
import chromadb

CHROMA_PATH      = "./chroma_db"
COLLECTION_NAME  = "docmind_collection"
EMBED_MODEL_NAME = "BAAI/bge-base-en-v1.5"   # upgraded from all-MiniLM-L6-v2
CHUNK_SIZE       = 400
CHUNK_OVERLAP    = 75

SUPPORTED_TYPES = {
    "pdf":  [".pdf"],
    "word": [".docx"],
}

_embed_model   = None
_chroma_client = None
_collection    = None


def get_embed_model():
    global _embed_model
    if _embed_model is None:
        print(f"[Ingest] Loading embedding model: {EMBED_MODEL_NAME}")
        _embed_model = SentenceTransformer(EMBED_MODEL_NAME)
    return _embed_model


def get_collection():
    global _chroma_client, _collection
    if _collection is None:
        _chroma_client = chromadb.EphemeralClient()
        _collection = _chroma_client.get_or_create_collection(
            name=COLLECTION_NAME,
            metadata={"hnsw:space": "cosine"},
        )
        print(f"[Ingest] ChromaDB in-memory collection ready: {COLLECTION_NAME}")
    return _collection


def get_file_type(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    for ftype, exts in SUPPORTED_TYPES.items():
        if ext in exts:
            return ftype
    return "unknown"


# ── Extractors ────────────────────────────────────────────────────

def extract_from_pdf(file_bytes: bytes) -> tuple[str, int]:
    doc = pymupdf.open(stream=file_bytes, filetype="pdf")
    pages_text = [page.get_text() for page in doc]
    full_text  = "\n\n".join(pages_text)
    num_pages  = len(doc)
    doc.close()
    return full_text, num_pages


def extract_from_word(file_bytes: bytes, filename: str) -> tuple[str, int]:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".doc":
        raise ValueError(
            "Old .doc format is not supported. "
            "Please open in Word → Save As → .docx, then re-upload."
        )
    doc   = DocxDocument(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            if para.style.name.startswith('Heading'):
                parts.append(f"\n## {text}\n")
            else:
                parts.append(text)
    for i, table in enumerate(doc.tables):
        parts.append(f"\n[Table {i+1}]")
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                parts.append(row_text)
    full_text = "\n".join(parts)
    num_pages = max(1, len(full_text.split()) // 300)
    return full_text, num_pages


# ── Chunking + Storage ────────────────────────────────────────────

def split_into_chunks(text: str) -> list[str]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_text(text)


def embed_and_store(chunks: list[str], doc_id: str, filename: str, file_type: str) -> int:
    model      = get_embed_model()
    collection = get_collection()

    existing = collection.get(where={"doc_id": doc_id})
    if existing["ids"]:
        collection.delete(ids=existing["ids"])

    print(f"[Ingest] Embedding {len(chunks)} chunks with {EMBED_MODEL_NAME}…")

    # BGE models work best with a query prefix for passage encoding
    passages = [f"passage: {chunk}" for chunk in chunks]
    embeddings = model.encode(passages, show_progress_bar=False, normalize_embeddings=True).tolist()

    ids       = [str(uuid.uuid4()) for _ in chunks]
    metadatas = [
        {"doc_id": doc_id, "filename": filename, "file_type": file_type, "chunk_index": i}
        for i in range(len(chunks))
    ]

    for start in range(0, len(chunks), 50):
        end = start + 50
        collection.add(
            ids=ids[start:end],
            embeddings=embeddings[start:end],
            documents=chunks[start:end],
            metadatas=metadatas[start:end],
        )

    print(f"[Ingest] Stored {len(chunks)} chunks.")
    return len(chunks)


# ── Main ──────────────────────────────────────────────────────────

def ingest_file(file_bytes: bytes, filename: str, api_key: str = "") -> dict:
    doc_id    = filename.replace(" ", "_").lower()
    file_type = get_file_type(filename)

    if file_type == "unknown":
        raise ValueError("Only PDF (.pdf) and Word (.docx) files are supported.")

    print(f"[Ingest] Starting: {filename} (type={file_type})")

    if file_type == "pdf":
        text, pages = extract_from_pdf(file_bytes)
    elif file_type == "word":
        text, pages = extract_from_word(file_bytes, filename)

    word_count = len(text.split())
    if word_count < 5:
        raise ValueError("File appears empty or has no extractable content.")

    chunks     = split_into_chunks(text)
    num_stored = embed_and_store(chunks, doc_id, filename, file_type)

    return {
        "doc_id":    doc_id,
        "filename":  filename,
        "file_type": file_type,
        "pages":     pages,
        "words":     word_count,
        "chunks":    num_stored,
    }